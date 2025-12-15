# app/docs_contractual/views.py
from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.db.models import Q
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from login.decorators import modulo_required
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.conf import settings
from dashboard.models import Contrato
from .models import HistorialGeneracion, PlantillaDocumento
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime, date
import tempfile
import re
import locale

# Configurar locale para fechas en español
try:
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'Spanish_Spain.1252')
    except:
        pass  # Si no se puede configurar, usar nombres manuales

# Diccionario para nombres de meses en español
MESES_ESPANOL = {
    'January': 'enero', 'February': 'febrero', 'March': 'marzo',
    'April': 'abril', 'May': 'mayo', 'June': 'junio',
    'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
    'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
}

# Código de entidad (mismo que dashboard)
codigo_ent = 727001372

def formatear_fecha_espanol(fecha):
    """Formatear fecha en español"""
    if not fecha:
        return 'N/A'
    
    try:
        # Formatear fecha en inglés primero
        fecha_str = fecha.strftime('%d de %B de %Y')
        
        # Reemplazar nombres de meses en inglés por español
        for ingles, espanol in MESES_ESPANOL.items():
            fecha_str = fecha_str.replace(ingles, espanol)
        
        return fecha_str
    except:
        return fecha.strftime('%d/%m/%Y') if fecha else 'N/A'

@login_required
@modulo_required('docs_contractual')
def docs_contractual_index(request):
    """Vista principal del módulo docs_contractual con filtros y paginación"""
    # Obtener parámetros de filtros
    año_filtro = request.GET.get('año', '2025')  # Por defecto 2025
    referencia_filtro = request.GET.get('referencia', '').strip()
    proveedor_filtro = request.GET.get('proveedor', '').strip()
    registros_filtro = request.GET.get('registros', '25')  # Por defecto 25 registros
    pagina_actual = request.GET.get('pagina', '1')  # Página actual
    
    try:
        año_filtro = int(año_filtro)
    except (ValueError, TypeError):
        año_filtro = 2025
    
    try:
        pagina_actual = int(pagina_actual)
        if pagina_actual < 1:
            pagina_actual = 1
    except (ValueError, TypeError):
        pagina_actual = 1
    
    # Query base
    queryset = Contrato.objects.filter(codigo_entidad=codigo_ent)
    
    # Aplicar filtro de año
    queryset = queryset.filter(fecha_de_firma__year=año_filtro)
    
    # Aplicar filtros de búsqueda si existen
    if referencia_filtro:
        queryset = queryset.filter(referencia_del_contrato__icontains=referencia_filtro)
    
    if proveedor_filtro:
        queryset = queryset.filter(proveedor_adjudicado__icontains=proveedor_filtro)
    
    # Aplicar ambos filtros con OR si se proporcionan ambos
    if referencia_filtro and proveedor_filtro:
        queryset = Contrato.objects.filter(
            codigo_entidad=codigo_ent,
            fecha_de_firma__year=año_filtro
        ).filter(
            Q(referencia_del_contrato__icontains=referencia_filtro) |
            Q(proveedor_adjudicado__icontains=proveedor_filtro)
        )
    
    # Ordenar por fecha más reciente
    queryset = queryset.order_by('-fecha_de_firma')
    
    # Contar total antes de aplicar limit
    total_contratos = queryset.count()
    
    # Variables de paginación
    tiene_paginacion = False
    total_paginas = 1
    pagina_siguiente = None
    pagina_anterior = None
    rango_paginas = []
    offset = 0
    
    # Aplicar paginación según registros seleccionados
    if registros_filtro != 'todos':
        try:
            registros_por_pagina = int(registros_filtro)
            
            # Calcular paginación
            total_paginas = max(1, (total_contratos + registros_por_pagina - 1) // registros_por_pagina)
            
            # Asegurar que la página actual no exceda el total
            if pagina_actual > total_paginas:
                pagina_actual = total_paginas
            
            # Calcular offset
            offset = (pagina_actual - 1) * registros_por_pagina
            
            # Obtener contratos para la página actual
            contratos = queryset[offset:offset + registros_por_pagina]
            
            # Configurar variables de paginación
            if total_paginas > 1:
                tiene_paginacion = True
                
                # Página anterior
                if pagina_actual > 1:
                    pagina_anterior = pagina_actual - 1
                
                # Página siguiente
                if pagina_actual < total_paginas:
                    pagina_siguiente = pagina_actual + 1
                
                # Rango de páginas a mostrar (máximo 5 páginas)
                inicio_rango = max(1, pagina_actual - 2)
                fin_rango = min(total_paginas, pagina_actual + 2)
                
                # Ajustar el rango si hay menos de 5 páginas
                if fin_rango - inicio_rango < 4:
                    if inicio_rango == 1:
                        fin_rango = min(total_paginas, inicio_rango + 4)
                    else:
                        inicio_rango = max(1, fin_rango - 4)
                
                rango_paginas = list(range(inicio_rango, fin_rango + 1))
            
        except (ValueError, TypeError):
            contratos = queryset[:25]
            registros_filtro = '25'
    else:
        contratos = queryset
    
    # Obtener años disponibles para el select
    años_disponibles = Contrato.objects.filter(
        codigo_entidad=codigo_ent,
        fecha_de_firma__isnull=False
    ).dates('fecha_de_firma', 'year', order='DESC')
    años_disponibles = [fecha.year for fecha in años_disponibles]
    
    # Asegurar que 2025 esté en la lista
    if 2025 not in años_disponibles:
        años_disponibles.insert(0, 2025)
        años_disponibles.sort(reverse=True)
    
    # Si no hay años disponibles, usar 2025
    if not años_disponibles:
        años_disponibles = [2025]
    
    # Calcular el rango de registros mostrados
    if registros_filtro != 'todos':
        registro_inicio = offset + 1 if total_contratos > 0 else 0
        registro_fin = min(offset + len(contratos), total_contratos)
    else:
        registro_inicio = 1 if total_contratos > 0 else 0
        registro_fin = total_contratos
    
    context = {
        'title': 'Docs Contractual',
        'contratos': contratos,
        'años_disponibles': años_disponibles,
        'año_actual': 2025,  # Año por defecto
        'año_seleccionado': año_filtro,
        'total_contratos': total_contratos,
        'referencia_filtro': referencia_filtro,
        'proveedor_filtro': proveedor_filtro,
        'registros_por_pagina': registros_filtro,
        # Variables de paginación
        'tiene_paginacion': tiene_paginacion,
        'pagina_actual': pagina_actual,
        'total_paginas': total_paginas,
        'pagina_anterior': pagina_anterior,
        'pagina_siguiente': pagina_siguiente,
        'rango_paginas': rango_paginas,
        'registro_inicio': registro_inicio,
        'registro_fin': registro_fin,
    }
    
    return render(request, 'docs_contractual/index.html', context)

@login_required 
def buscar_contratos(request):
    """Vista AJAX para búsqueda de contratos (mantener para compatibilidad pero ahora integrada en index)"""
    if request.method == 'GET':
        referencia = request.GET.get('referencia', '').strip()
        proveedor = request.GET.get('proveedor', '').strip()
        año = request.GET.get('año', '2025')
        
        # Si no hay parámetros de búsqueda, retornar contratos del año
        try:
            año = int(año)
        except (ValueError, TypeError):
            año = 2025
        
        # Query base
        queryset = Contrato.objects.filter(codigo_entidad=codigo_ent, fecha_de_firma__year=año)
        
        # Aplicar filtros de búsqueda si existen
        if referencia:
            queryset = queryset.filter(referencia_del_contrato__icontains=referencia)
        
        if proveedor:
            queryset = queryset.filter(proveedor_adjudicado__icontains=proveedor)
        
        # Si se proporcionaron ambos filtros, usar OR
        if referencia and proveedor:
            queryset = Contrato.objects.filter(
                codigo_entidad=codigo_ent,
                fecha_de_firma__year=año
            ).filter(
                Q(referencia_del_contrato__icontains=referencia) |
                Q(proveedor_adjudicado__icontains=proveedor)
            )
        
        # Limitar resultados a 50 para mejor rendimiento
        contratos = queryset.order_by('-fecha_de_firma')[:50]
        
        # Preparar datos para JSON
        resultados = []
        for contrato in contratos:
            resultados.append({
                'id': contrato.id,
                'referencia_del_contrato': contrato.referencia_del_contrato or '',
                'objeto_del_contrato': contrato.descripcion_del_proceso or '',
                'estado_contrato': contrato.estado_contrato or '',
                'proveedor_adjudicado': contrato.proveedor_adjudicado or '',
                # Datos adicionales para el documento
                'fecha_de_firma': contrato.fecha_de_firma.strftime('%Y-%m-%d') if contrato.fecha_de_firma else '',
                'tipo_de_contrato': contrato.tipo_de_contrato or '',
                'tipodocproveedor': contrato.tipodocproveedor or '',
                'documento_proveedor': contrato.documento_proveedor or '',
                'valor_del_contrato': str(contrato.valor_del_contrato) if contrato.valor_del_contrato else '0',
                'duracion_del_contrato': getattr(contrato, 'duracion_del_contrato', '') or ''
            })
        
        return JsonResponse({
            'success': True,
            'contratos': resultados,
            'total': len(resultados)
        })
    
    return JsonResponse({'success': False, 'error': 'Método no permitido'})

def aplicar_formato_arial10(paragraph):
    """Aplicar formato Arial 10 a un párrafo"""
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)

def reemplazar_variables_en_docx(doc_path, variables):
    """
    Función para reemplazar variables en un documento Word con formato Arial 10
    
    Args:
        doc_path: Ruta al archivo .docx plantilla
        variables: Diccionario con variables a reemplazar
    
    Returns:
        Document object con las variables reemplazadas
    """
    try:
        # Cargar el documento plantilla
        doc = Document(doc_path)
        
        # Reemplazar variables en párrafos
        for paragraph in doc.paragraphs:
            for variable, valor in variables.items():
                if variable in paragraph.text:
                    paragraph.text = paragraph.text.replace(variable, str(valor))
                    aplicar_formato_arial10(paragraph)
        
        # Reemplazar variables en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for variable, valor in variables.items():
                            if variable in paragraph.text:
                                paragraph.text = paragraph.text.replace(variable, str(valor))
                                aplicar_formato_arial10(paragraph)
        
        return doc
        
    except Exception as e:
        print(f"Error al reemplazar variables: {e}")
        return None

@login_required
@require_POST
def generar_documento(request):
    """Vista para generar documento Word con información del contrato usando plantilla"""
    try:
        # Obtener datos del POST (form data o JSON)
        if request.content_type == 'application/json':
            data = json.loads(request.body)
            contrato_id = data.get('contrato_id')
        else:
            contrato_id = request.POST.get('contrato_id')
        
        if not contrato_id:
            return JsonResponse({'success': False, 'error': 'ID de contrato no proporcionado'})
        
        # Obtener el contrato
        contrato = get_object_or_404(Contrato, id=contrato_id, codigo_entidad=codigo_ent)
        
        # Verificar si existe plantilla personalizada
        plantilla_path = os.path.join(
            settings.BASE_DIR, 
            'docs_contractual', 
            'templates', 
            'documents', 
            'plantilla_designacion.docx'
        )
        
        if os.path.exists(plantilla_path):
            # Usar plantilla personalizada
            tipodoc_valor = 'NIT' if (contrato.tipodocproveedor or '').upper() in ['NO DEFINIDO', 'N/A', ''] else (contrato.tipodocproveedor or 'N/A').upper()
            
            variables = {
                '{{w_fecha}}': formatear_fecha_espanol(contrato.fecha_de_firma),
                '{{w_contrato}}': (contrato.referencia_del_contrato or 'N/A').upper(),
                '{{w_tipo}}': (contrato.tipo_de_contrato or 'N/A').upper(),
                '{{w_contratista}}': (contrato.proveedor_adjudicado or 'N/A').upper(),
                '{{w_tipodoc}}': tipodoc_valor,
                '{{w_id}}': (contrato.documento_proveedor or 'N/A').upper(),
                '{{w_objeto}}': (contrato.descripcion_del_proceso or 'N/A').upper(),
                '{{w_valor}}': f"${contrato.valor_del_contrato:,.0f}" if contrato.valor_del_contrato else 'N/A',
                '{{w_plazo}}': (getattr(contrato, 'duracion_del_contrato', '') or 'N/A').upper() if getattr(contrato, 'duracion_del_contrato', '') else 'N/A'
            }
            
            # Generar documento usando plantilla
            doc = reemplazar_variables_en_docx(plantilla_path, variables)
            if not doc:
                raise Exception("Error al procesar la plantilla")
            
        else:
            # Generar documento por defecto si no hay plantilla
            doc = Document()
            
            # Configuración de la página
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Título del documento
            titulo = doc.add_paragraph()
            titulo_run = titulo.add_run('DESIGNACIÓN DE SUPERVISOR DEL CONTRATO')
            titulo_run.bold = True
            titulo_run.font.name = 'Arial'
            titulo_run.font.size = Pt(12)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph('')
            
            # Formatear datos del contrato
            tipodoc_valor = 'NIT' if (contrato.tipodocproveedor or '').upper() in ['NO DEFINIDO', 'N/A', ''] else (contrato.tipodocproveedor or 'N/A').upper()
            
            # Contenido del documento
            contenido_parrafos = [
                f"En mi calidad de Ordenador del Gasto, procedo a designar supervisor para el siguiente contrato:",
                f"",
                f"CONTRATO: {(contrato.referencia_del_contrato or 'N/A').upper()}",
                f"TIPO: {(contrato.tipo_de_contrato or 'N/A').upper()}",
                f"CONTRATISTA: {(contrato.proveedor_adjudicado or 'N/A').upper()}",
                f"{tipodoc_valor}: {(contrato.documento_proveedor or 'N/A').upper()}",
                f"OBJETO: {(contrato.descripcion_del_proceso or 'N/A').upper()}",
                f"VALOR: ${contrato.valor_del_contrato:,.0f}" if contrato.valor_del_contrato else "VALOR: N/A",
                f"FECHA DE FIRMA: {formatear_fecha_espanol(contrato.fecha_de_firma)}",
                f"",
                f"La presente designación se realiza de conformidad con la normatividad vigente."
            ]
            
            # Agregar párrafos con formato
            for texto in contenido_parrafos:
                p = doc.add_paragraph(texto)
                aplicar_formato_arial10(p)
            
            # Espacio para firmas
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            # Firmas
            linea1 = doc.add_paragraph('_' * 40)
            aplicar_formato_arial10(linea1)
            
            ordenador = doc.add_paragraph('ORDENADOR DEL GASTO')
            aplicar_formato_arial10(ordenador)
            
            doc.add_paragraph('')
            
            linea2 = doc.add_paragraph('_' * 40)
            aplicar_formato_arial10(linea2)
            
            supervisor = doc.add_paragraph('SUPERVISOR DESIGNADO')
            aplicar_formato_arial10(supervisor)
        
        # Generar nombre del archivo
        referencia_limpia = contrato.referencia_del_contrato.replace('/', '_').replace(' ', '_') if contrato.referencia_del_contrato else 'contrato'
        nombre_archivo = f"Designacion_Contrato_{referencia_limpia}.docx"
        
        # Crear respuesta HTTP con el documento
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
        
        # Guardar el documento en la respuesta
        doc.save(response)
        
        return response
        
    except Exception as e:
        print(f"Error al generar documento: {e}")
        return JsonResponse({
            'success': False, 
            'error': f'Error al generar el documento: {str(e)}'
        })

@login_required
def historial_documentos(request):
    """Vista para mostrar historial de documentos generados"""
    historial = HistorialGeneracion.objects.all().order_by('-fecha_generacion')[:100]
    
    return render(request, 'docs_contractual/historial.html', {
        'historial': historial,
        'title': 'Historial de Documentos'
    })

@login_required
def gestionar_plantillas(request):
    """Vista para gestionar plantillas de documentos"""
    # Verificar si existe la plantilla por defecto
    plantilla_path = os.path.join(
        settings.BASE_DIR, 
        'docs_contractual', 
        'templates', 
        'documents', 
        'plantilla_designacion.docx'
    )
    
    plantilla_existe = os.path.exists(plantilla_path)
    
    context = {
        'title': 'Gestión de Plantillas',
        'plantilla_existe': plantilla_existe,
        'plantilla_path': 'docs_contractual/templates/documents/plantilla_designacion.docx',
        'variables_disponibles': [
            '{{w_fecha}} - Fecha de firma del contrato (en español)',
            '{{w_contrato}} - Referencia del contrato (MAYÚSCULAS)',
            '{{w_tipo}} - Tipo de contrato (MAYÚSCULAS)',
            '{{w_contratista}} - Nombre del proveedor (MAYÚSCULAS)',
            '{{w_tipodoc}} - Tipo de documento ("No Definido" → "NIT")',
            '{{w_id}} - Número de documento del proveedor (MAYÚSCULAS)',
            '{{w_objeto}} - Objeto del contrato (MAYÚSCULAS)',
            '{{w_valor}} - Valor del contrato (formato moneda)',
            '{{w_plazo}} - Duración del contrato (MAYÚSCULAS)'
        ]
    }
    
    return render(request, 'docs_contractual/plantillas.html', context)

@login_required
def generar_designacion_get(request, contrato_id):
    """Vista GET para generar documento de designación directamente desde el listado"""
    try:
        # Obtener el contrato
        contrato = get_object_or_404(Contrato, id=contrato_id, codigo_entidad=codigo_ent)
        
        # Verificar si existe plantilla personalizada
        plantilla_path = os.path.join(
            settings.BASE_DIR, 
            'docs_contractual', 
            'templates', 
            'documents', 
            'plantilla_designacion.docx'
        )
        
        if os.path.exists(plantilla_path):
            # Usar plantilla personalizada
            tipodoc_valor = 'NIT' if (contrato.tipodocproveedor or '').upper() in ['NO DEFINIDO', 'N/A', ''] else (contrato.tipodocproveedor or 'N/A').upper()
            
            variables = {
                '{{w_fecha}}': formatear_fecha_espanol(contrato.fecha_de_firma),
                '{{w_contrato}}': (contrato.referencia_del_contrato or 'N/A').upper(),
                '{{w_tipo}}': (contrato.tipo_de_contrato or 'N/A').upper(),
                '{{w_contratista}}': (contrato.proveedor_adjudicado or 'N/A').upper(),
                '{{w_tipodoc}}': tipodoc_valor,
                '{{w_id}}': (contrato.documento_proveedor or 'N/A').upper(),
                '{{w_objeto}}': (contrato.descripcion_del_proceso or 'N/A').upper(),
                '{{w_valor}}': f"${contrato.valor_del_contrato:,.0f}" if contrato.valor_del_contrato else 'N/A',
                '{{w_plazo}}': (getattr(contrato, 'duracion_del_contrato', '') or 'N/A').upper() if getattr(contrato, 'duracion_del_contrato', '') else 'N/A'
            }
            
            # Generar documento usando plantilla
            doc = reemplazar_variables_en_docx(plantilla_path, variables)
            if not doc:
                raise Exception("Error al procesar la plantilla")
            
        else:
            # Generar documento por defecto si no hay plantilla
            doc = Document()
            
            # Configuración de la página
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Título del documento
            titulo = doc.add_paragraph()
            titulo_run = titulo.add_run('DESIGNACIÓN DE SUPERVISOR DEL CONTRATO')
            titulo_run.bold = True
            titulo_run.font.name = 'Arial'
            titulo_run.font.size = Pt(12)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph('')
            
            # Formatear datos del contrato
            tipodoc_valor = 'NIT' if (contrato.tipodocproveedor or '').upper() in ['NO DEFINIDO', 'N/A', ''] else (contrato.tipodocproveedor or 'N/A').upper()
            
            # Contenido del documento
            contenido_parrafos = [
                f"En mi calidad de Ordenador del Gasto, procedo a designar supervisor para el siguiente contrato:",
                f"",
                f"CONTRATO: {(contrato.referencia_del_contrato or 'N/A').upper()}",
                f"TIPO: {(contrato.tipo_de_contrato or 'N/A').upper()}",
                f"CONTRATISTA: {(contrato.proveedor_adjudicado or 'N/A').upper()}",
                f"{tipodoc_valor}: {(contrato.documento_proveedor or 'N/A').upper()}",
                f"OBJETO: {(contrato.descripcion_del_proceso or 'N/A').upper()}",
                f"VALOR: ${contrato.valor_del_contrato:,.0f}" if contrato.valor_del_contrato else "VALOR: N/A",
                f"FECHA DE FIRMA: {formatear_fecha_espanol(contrato.fecha_de_firma)}",
                f"",
                f"La presente designación se realiza de conformidad con la normatividad vigente."
            ]
            
            # Agregar párrafos con formato
            for texto in contenido_parrafos:
                p = doc.add_paragraph(texto)
                aplicar_formato_arial10(p)
            
            # Espacio para firmas
            doc.add_paragraph('')
            doc.add_paragraph('')
            
            # Firmas
            linea1 = doc.add_paragraph('_' * 40)
            aplicar_formato_arial10(linea1)
            
            ordenador = doc.add_paragraph('ORDENADOR DEL GASTO')
            aplicar_formato_arial10(ordenador)
            
            doc.add_paragraph('')
            
            linea2 = doc.add_paragraph('_' * 40)
            aplicar_formato_arial10(linea2)
            
            supervisor = doc.add_paragraph('SUPERVISOR DESIGNADO')
            aplicar_formato_arial10(supervisor)
        
        # Generar nombre del archivo
        fecha_actual = datetime.now().strftime('%Y%m%d_%H%M%S')
        referencia_limpia = contrato.referencia_del_contrato.replace('/', '_').replace(' ', '_') if contrato.referencia_del_contrato else 'contrato'
        nombre_archivo = f"Designacion_{referencia_limpia}_{fecha_actual}.docx"
        
        # Crear respuesta HTTP con el documento
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
        
        # Guardar el documento en la respuesta
        doc.save(response)
        
        return response
        
    except Exception as e:
        print(f"Error al generar documento: {e}")
        return HttpResponse(f'Error al generar el documento: {str(e)}', status=500)

@login_required
def preview_variables(request, contrato_id):
    """Vista para preview de variables de un contrato específico con formato correcto"""
    try:
        contrato = get_object_or_404(Contrato, id=contrato_id, codigo_entidad=codigo_ent)
        
        # Aplicar el mismo formato que se usa en la generación real
        tipodoc_valor = 'NIT' if (contrato.tipodocproveedor or '').upper() in ['NO DEFINIDO', 'N/A', ''] else (contrato.tipodocproveedor or 'N/A').upper()
        
        variables = {
            'w_fecha': formatear_fecha_espanol(contrato.fecha_de_firma),
            'w_contrato': (contrato.referencia_del_contrato or 'N/A').upper(),
            'w_tipo': (contrato.tipo_de_contrato or 'N/A').upper(),
            'w_contratista': (contrato.proveedor_adjudicado or 'N/A').upper(),
            'w_tipodoc': tipodoc_valor,
            'w_id': (contrato.documento_proveedor or 'N/A').upper(),
            'w_objeto': (contrato.descripcion_del_proceso or 'N/A').upper(),
            'w_valor': f"${contrato.valor_del_contrato:,.0f}" if contrato.valor_del_contrato else 'N/A',
            'w_plazo': (getattr(contrato, 'duracion_del_contrato', '') or 'N/A').upper() if getattr(contrato, 'duracion_del_contrato', '') else 'N/A'
        }
        
        return JsonResponse({
            'success': True,
            'variables': variables,
            'contrato_info': {
                'referencia': contrato.referencia_del_contrato,
                'proveedor': contrato.proveedor_adjudicado
            },
            'formato_info': {
                'fuente': 'Arial 10',
                'fecha_idioma': 'Español',
                'variables_mayusculas': 'Sí (excepto fecha)',
                'tipodoc_conversion': '"No Definido" se convierte a "NIT"'
            }
        })
        
    except Contrato.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Contrato no encontrado'
        })