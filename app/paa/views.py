import os
import re
import json
import logging
import tempfile
import time
from datetime import datetime
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.conf import settings
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
from pathlib import Path

# Configurar logger
logger = logging.getLogger(__name__)

# Configurar Gemini API
API_KEY = "AIzaSyB-5c2Co7MuxWZnleGv_rjdSP7lrfOmRnM"
genai.configure(api_key=API_KEY)


@login_required
def index(request):
    """Vista principal para generar certificado PAA"""
    return render(request, 'paa/index.html')


@login_required
@csrf_exempt
def generar_certificado(request):
    """Procesa el estudio previo y genera el certificado PAA - VERSIÓN SIMPLIFICADA"""
    print("\n" + "="*60)
    print("INICIANDO GENERACIÓN DE CERTIFICADO PAA (MODO SIMPLIFICADO)")
    print("="*60)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Método no permitido'}, status=405)
    
    if 'estudio_previo' not in request.FILES:
        return JsonResponse({'error': 'No se ha cargado ningún archivo'}, status=400)
    
    try:
        # Obtener el archivo cargado
        estudio_previo = request.FILES['estudio_previo']
        print(f"✓ Archivo recibido: {estudio_previo.name} ({estudio_previo.size} bytes)")
        
        # Validar que sea un archivo .docx
        if not estudio_previo.name.endswith('.docx'):
            return JsonResponse({'error': 'El archivo debe ser un documento Word (.docx)'}, status=400)
        
        # Extraer texto del documento
        try:
            print("✓ Extrayendo texto del documento...")
            texto_final = extraer_texto_de_docx(estudio_previo)
            
            # Buscar códigos UNSPSC en el texto extraído (para debugging)
            import re
            codigos_encontrados = re.findall(r'\b\d{8}\b', texto_final)
            print(f"✓ Códigos de 8 dígitos encontrados en el texto: {codigos_encontrados[:10]}")
            
        except Exception as e:
            print(f"✗ ERROR al extraer texto: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'Error al extraer texto del documento: {str(e)}'}, status=500)
        
        # Extraer información usando Gemini
        try:
            print("✓ Procesando texto con Gemini...")
            info_extraida = extraer_informacion_con_gemini_file(texto_final)
        except Exception as e:
            print(f"✗ ERROR en Gemini: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'Error al procesar con Gemini: {str(e)}'}, status=500)
        
        # Información fija del certificado (hardcodeada)
        nombre_completo = 'ANDRÉS OCAMPO CASTAÑO'
        cargo = 'PROFESIONAL UNIVERSITARIO SISTEMAS DE INFORMACIÓN'
        genero = 'EL'
        
        # Convertir lista de códigos UNSPSC a string
        codigos_unspsc = info_extraida.get('CODIGOS_UNSPSC', [])
        if isinstance(codigos_unspsc, list):
            codigos_str = ', '.join(codigos_unspsc)
        else:
            codigos_str = str(codigos_unspsc)
        
        # Preparar datos para el certificado (mapeo de campos)
        datos = {
            'w_nom_funcionario': nombre_completo,
            'w_gen': genero,
            'w_cargo': cargo,
            'w_anno': str(datetime.now().year),
            'w_codigos': codigos_str,
            'w_codigo_paa': info_extraida.get('CODIGO_PAA', ''),
            'w_objeto': info_extraida.get('OBJETO', ''),
            'w_valor': info_extraida.get('VALOR', ''),
            'w_plazo': info_extraida.get('PLAZO', ''),
            'w_fecha': info_extraida.get('FECHA', '')
        }
        
        print("\n" + "="*60)
        print("DATOS FINALES PARA EL CERTIFICADO:")
        print("="*60)
        for key, value in datos.items():
            print(f"{key}: {str(value)[:100]}..." if len(str(value)) > 100 else f"{key}: {value}")
        print("="*60 + "\n")
        
        # Generar el certificado Word
        try:
            print("✓ Generando documento PAA...")
            certificado = generar_documento_paa(datos)
            print("✓ Documento generado exitosamente")
        except FileNotFoundError as e:
            print(f"✗ ERROR: Plantilla no encontrada - {str(e)}")
            return JsonResponse({'error': 'No se encontró la plantilla PAA.'}, status=500)
        except Exception as e:
            print(f"✗ ERROR al generar documento: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'Error al generar el certificado: {str(e)}'}, status=500)
        
        # Preparar respuesta Word
        print("✓ Preparando respuesta HTTP...")
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Certificado_PAA_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
        certificado.save(response)
        
        print("✓ CERTIFICADO GENERADO EXITOSAMENTE")
        print("="*60 + "\n")
        return response
        
    except Exception as e:
        # Log detallado del error
        print("\n" + "="*60)
        print("✗✗✗ ERROR GENERAL ✗✗✗")
        print("="*60)
        logger.error(f"Error en generar_certificado: {str(e)}", exc_info=True)
        import traceback
        error_detail = traceback.format_exc()
        print(f"{error_detail}")
        print("="*60 + "\n")
        return JsonResponse({'error': f'Error al procesar el documento: {str(e)}'}, status=500)


def extraer_texto_de_docx(archivo_word):
    """
    Extrae todo el texto de un archivo .docx (párrafos y tablas).
    
    Args:
        archivo_word: Archivo Django UploadedFile (.docx)
        
    Returns:
        str: Texto completo del documento
    """
    print("✓ Extrayendo texto del documento Word...")
    if hasattr(archivo_word, 'seek'):
        archivo_word.seek(0)

    document = Document(archivo_word)
    texto_completo = []

    # Extraer párrafos
    for parrafo in document.paragraphs:
        texto_completo.append(parrafo.text)

    # Extraer tablas
    for tabla in document.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                texto_completo.append(celda.text)

    texto_final = '\n'.join(texto_completo)
    print(f"✓ Texto extraído: {len(texto_final)} caracteres")
    
    return texto_final


def extraer_informacion_con_gemini_file(texto_final):
    """
    Extrae información estructurada del estudio previo usando Gemini.
    Utiliza un prompt preciso para obtener datos en formato JSON.
    
    Args:
        texto_final: Texto completo extraído del documento Word
        
    Returns:
        dict: Información extraída con claves: CODIGOS_UNSPSC, CODIGO_PAA, OBJETO, VALOR, PLAZO, FECHA
    """
    print("✓ Inicializando modelo gemini-2.0-flash-exp...")
    model = genai.GenerativeModel('gemini-2.0-flash-exp')
    
    try:
        # 1. Construir prompt estructurado
        prompt = f"""
Eres un asistente experto en la extracción de datos de documentos de contratación pública, específicamente "Estudios Previos". Tu objetivo es extraer información clave para rellenar automáticamente un formato de certificado PAA.

INSTRUCCIÓN CLAVE:
1. Debes analizar el texto proporcionado a continuación y extraer los SEIS (6) campos solicitados.
2. Tu respuesta debe ser **ESTRICTAMENTE** un objeto JSON válido y nada más. No incluyas preámbulos, explicaciones ni código Markdown adicional (como ```json```).
3. Utiliza **EXACTAMENTE** las siguientes claves JSON, ya que están mapeadas a la plantilla del certificado:

CLAVES JSON REQUERIDAS y DEFINICIONES DE VALOR:

1. **"CODIGOS_UNSPSC"**: Busca en el texto todos el código o los códigos UNSPSC que existan en todo el documento.

2. **"CODIGO_PAA"**: El código PAA del proyecto. Busca un código que comience con "2025-" seguido de números (ejemplo: "2025-123456"). Si no se encuentra, deja vacío: "".

3. **"OBJETO"**: El texto completo y limpio del objeto del contrato (la descripción de lo que se va a contratar). Debe estar en MAYÚSCULAS.

4. **"VALOR"**: El valor total del contrato. Debe contener la cifra en letras seguida del valor numérico exacto entre paréntesis (ej. "CIENTO ONCE MILLONES TRESCIENTOS SESENTA Y CINCO MIL DOSCIENTOS SETENTA Y TRES PESOS M/L ($ 111.365.273)").

5. **"PLAZO"**: La duración exacta del contrato tal como se especifica en la sección de Plazo (ej. "SIETE (7) MESES CONTADOS A PARTIR DE LA INDICACIÓN DEL ACTA DE INICIO").

6. **"FECHA"**: Extrae la fecha actual en español (ej. "el doce (12) días del mes de diciembre de 2025"). Si no se encuentra una fecha explícita, genera la fecha actual en ese formato.

EJEMPLO DE RESPUESTA ESPERADA:
{{
  "CODIGOS_UNSPSC": ["80111600", "70101100"],
  "CODIGO_PAA": "2025-123456",
  "OBJETO": "PRESTACIÓN DE SERVICIOS PROFESIONALES...",
  "VALOR": "CIENTO ONCE MILLONES... ($ 111.365.273)",
  "PLAZO": "SIETE (7) MESES...",
  "FECHA": "el doce (12) días del mes de diciembre de 2025"
}}

IMPORTANTE: Busca cuidadosamente en TODO el texto los códigos UNSPSC. Pueden estar en secciones como "Clasificación UNSPSC", "Códigos", tablas de productos/servicios, o cualquier lugar donde aparezcan números de 8 dígitos relacionados con clasificación de productos.

TEXTO DEL DOCUMENTO A PROCESAR:
---
{texto_final}
---

RESPUESTA (SOLO JSON):
"""
        
        # 3. Enviar a Gemini
        print("✓ Enviando prompt a Gemini...")
        response = model.generate_content(prompt)
        
        if not response or not response.text:
            raise Exception("Gemini no devolvió respuesta válida")
        
        print(f"✓ Respuesta recibida de Gemini ({len(response.text)} caracteres)")
        texto_respuesta = response.text.strip()
        
        # Imprimir respuesta completa de Gemini para revisión
        print("\n" + "="*80)
        print("RESPUESTA COMPLETA DE GEMINI:")
        print("="*80)
        print(texto_respuesta)
        print("="*80 + "\n")
        
        # 4. Limpiar respuesta (eliminar markdown si existe)
        texto_respuesta = re.sub(r'```json\s*', '', texto_respuesta)
        texto_respuesta = re.sub(r'```\s*', '', texto_respuesta)
        texto_respuesta = texto_respuesta.strip()
        
        # 5. Parsear JSON
        print("✓ Parseando respuesta JSON...")
        json_data = json.loads(texto_respuesta)
        
        # 6. Validar estructura
        campos_requeridos = ['CODIGOS_UNSPSC', 'CODIGO_PAA', 'OBJETO', 'VALOR', 'PLAZO', 'FECHA']
        for campo in campos_requeridos:
            if campo not in json_data:
                print(f"⚠️ Campo faltante: {campo}")
                json_data[campo] = [] if campo == 'CODIGOS_UNSPSC' else ''
        
        # 7. Validar y extraer códigos UNSPSC con respaldo
        codigos = json_data.get('CODIGOS_UNSPSC', [])
        
        # Si no se encontraron códigos o está vacío, intentar extracción directa con regex
        if not codigos or (isinstance(codigos, list) and len(codigos) == 0):
            print("⚠️ Gemini no encontró códigos UNSPSC. Intentando extracción directa con regex...")
            codigos_regex = re.findall(r'\b\d{8}\b', texto_final)
            
            # Filtrar códigos que probablemente sean UNSPSC (buscar contexto)
            codigos_validos = []
            texto_lower = texto_final.lower()
            for codigo in codigos_regex:
                idx = texto_final.find(codigo)
                if idx != -1:
                    contexto = texto_final[max(0, idx-50):min(len(texto_final), idx+50)].lower()
                    if 'unspsc' in contexto or 'código' in contexto or 'clasificación' in contexto:
                        if codigo not in codigos_validos:
                            codigos_validos.append(codigo)
            
            if codigos_validos:
                json_data['CODIGOS_UNSPSC'] = codigos_validos
                print(f"✓ Códigos UNSPSC extraídos con regex: {codigos_validos}")
            else:
                print(f"⚠️ No se encontraron códigos UNSPSC con contexto. Todos los números de 8 dígitos: {codigos_regex[:5]}")
        
        # Mostrar códigos finales
        codigos = json_data.get('CODIGOS_UNSPSC', [])
        if isinstance(codigos, list):
            codigos_str = ', '.join(codigos)
        else:
            codigos_str = str(codigos)
        print(f"\n✓ CÓDIGOS UNSPSC FINALES: {codigos_str}\n")
        
        return json_data
        
    except json.JSONDecodeError as e:
        print(f"✗ Error parseando JSON: {str(e)}")
        print(f"Respuesta recibida: {texto_respuesta[:500]}...")
        # Retornar estructura por defecto
        return {
            'CODIGOS_UNSPSC': [],
            'CODIGO_PAA': '',
            'OBJETO': 'NO IDENTIFICADO',
            'VALOR': '0',
            'PLAZO': 'NO ESPECIFICADO',
            'FECHA': generar_fecha_en_letras()
        }
    except Exception as e:
        print(f"✗ Error al procesar con Gemini: {str(e)}")
        import traceback
        traceback.print_exc()
        raise


def generar_fecha_en_letras():
    """Genera la fecha actual en formato de letras en español"""
    meses = {
        1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
        5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
        9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
    }
    
    hoy = datetime.now()
    dia = hoy.day
    mes = meses[hoy.month]
    año = hoy.year
    
    # Convertir número a letras (simplificado para días 1-31)
    numeros = {
        1: 'uno', 2: 'dos', 3: 'tres', 4: 'cuatro', 5: 'cinco',
        6: 'seis', 7: 'siete', 8: 'ocho', 9: 'nueve', 10: 'diez',
        11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
        16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve',
        20: 'veinte', 21: 'veintiuno', 22: 'veintidós', 23: 'veintitrés',
        24: 'veinticuatro', 25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete',
        28: 'veintiocho', 29: 'veintinueve', 30: 'treinta', 31: 'treinta y uno'
    }
    
    dia_letras = numeros.get(dia, str(dia))
    
    return f"el {dia_letras} ({dia}) días del mes de {mes} de {año}"


def generar_documento_paa(datos):
    """Genera el documento PAA reemplazando los campos en la plantilla"""
    
    # Ruta a la plantilla
    base_dir = Path(__file__).resolve().parent
    plantilla_path = base_dir / 'templates' / 'paa' / 'plantilla_paa.docx'
    
    if not plantilla_path.exists():
        raise FileNotFoundError(f'No se encontró la plantilla en: {plantilla_path}')
    
    # Cargar la plantilla
    doc = Document(str(plantilla_path))
    
    # Reemplazar en párrafos
    for para in doc.paragraphs:
        for key, value in datos.items():
            placeholder = '{{' + key + '}}'
            if placeholder in para.text:
                # Reemplazar manteniendo el formato
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in datos.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in para.text:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
    
    # Reemplazar en encabezados y pies de página
    for section in doc.sections:
        # Encabezado
        header = section.header
        for para in header.paragraphs:
            for key, value in datos.items():
                placeholder = '{{' + key + '}}'
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
        
        # Pie de página
        footer = section.footer
        for para in footer.paragraphs:
            for key, value in datos.items():
                placeholder = '{{' + key + '}}'
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
    
    return doc


