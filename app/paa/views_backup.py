import os
import re
import logging
import tempfile
from datetime import datetime
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.conf import settings
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
from pathlib import Path
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("⚠️ docx2pdf no está instalado. La conversión a PDF no estará disponible.")

# Configurar logger
logger = logging.getLogger(__name__)


@login_required
def index(request):
    """Vista principal para generar certificado PAA"""
    return render(request, 'paa/index.html')


@login_required
@csrf_exempt
def generar_certificado(request):
    """Procesa el estudio previo y genera el certificado PAA"""
    print("\n" + "="*60)
    print("INICIANDO GENERACIÓN DE CERTIFICADO PAA")
    print("="*60)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Método no permitido'}, status=405)
    
    if 'estudio_previo' not in request.FILES:
        return JsonResponse({'error': 'No se ha cargado ningún archivo'}, status=400)
    
    try:
        print("Archivo recibido")
        # Obtener el archivo cargado
        estudio_previo = request.FILES['estudio_previo']
        
        # Validar que sea un archivo .docx
        if not estudio_previo.name.endswith('.docx'):
            return JsonResponse({'error': 'El archivo debe ser un documento Word (.docx)'}, status=400)
        
        # Leer el contenido del estudio previo
        try:
            print("Leyendo documento Word...")
            doc_estudio = Document(estudio_previo)
            
            # Extraer texto manteniendo el orden y contexto
            contenido = []
            
            # Agregar párrafos con marcador (con límite de caracteres)
            MAX_CHARS_PER_PARAGRAPH = 5000  # Límite por párrafo
            for para in doc_estudio.paragraphs:
                if para.text.strip():  # Solo si tiene contenido
                    texto_para = para.text[:MAX_CHARS_PER_PARAGRAPH]  # Limitar tamaño
                    contenido.append(texto_para)
            
            # Agregar tablas con mejor formato y contexto (con límite)
            MAX_CHARS_PER_CELL = 500  # Límite por celda
            MAX_ROWS_PER_TABLE = 100  # Límite de filas por tabla
            for idx, tabla in enumerate(doc_estudio.tables):
                contenido.append(f"\n--- TABLA {idx + 1} ---")
                for row_idx, fila in enumerate(tabla.rows):
                    if row_idx >= MAX_ROWS_PER_TABLE:
                        contenido.append("... (tabla truncada por tamaño)")
                        break
                    fila_texto = []
                    for celda in fila.cells:
                        # Limpiar texto de celda (quitar saltos de línea internos)
                        texto_limpio = celda.text.strip().replace('\n', ' ')[:MAX_CHARS_PER_CELL]
                        if texto_limpio:
                            fila_texto.append(texto_limpio)
                    if fila_texto:
                        contenido.append(" | ".join(fila_texto))
                contenido.append("--- FIN TABLA ---\n")
            
            # Combinar todo manteniendo el contexto
            texto_estudio = "\n".join(contenido)
            
            # LÍMITE CRÍTICO: Gemini tiene límite de ~100,000 caracteres (~30k tokens)
            MAX_CHARS_GEMINI = 80000  # Límite seguro para Gemini
            if len(texto_estudio) > MAX_CHARS_GEMINI:
                print(f"⚠️ Documento muy grande ({len(texto_estudio)} caracteres), truncando a {MAX_CHARS_GEMINI}")
                texto_estudio = texto_estudio[:MAX_CHARS_GEMINI]
                texto_estudio += "\n\n[DOCUMENTO TRUNCADO POR TAMAÑO]"
            
            print(f"✓ Documento leído: {len(texto_estudio)} caracteres TOTALES")
            
            # Debug: Buscar si existen códigos UNSPSC en el texto
            tiene_unspsc = 'UNSPSC' in texto_estudio.upper()
            print(f"✓ ¿Contiene la palabra 'UNSPSC'?: {tiene_unspsc}")
            
            # Buscar números de 8 dígitos (posibles códigos)
            import re
            codigos_8_digitos = re.findall(r'\b\d{8}\b', texto_estudio)
            if codigos_8_digitos:
                print(f"✓ Números de 8 dígitos encontrados: {codigos_8_digitos}")
            
            # Buscar códigos con formato (puntos, guiones, espacios)
            codigos_formateados = re.findall(r'\b\d{2}[\.\-\s]?\d{2}[\.\-\s]?\d{2}[\.\-\s]?\d{2}\b', texto_estudio)
            if codigos_formateados:
                print(f"✓ Códigos con formato encontrados: {codigos_formateados}")
            
            # Mostrar solo una muestra del texto (no todo el documento)
            muestra_inicio = texto_estudio[:1000]
            muestra_fin = texto_estudio[-500:] if len(texto_estudio) > 1500 else ""
            print(f"✓ Muestra del inicio del documento:\n{muestra_inicio}...")
            if muestra_fin:
                print(f"✓ Muestra del final del documento:\n...{muestra_fin}")
            print(f"✓ Enviando documento a Gemini ({len(texto_estudio)} caracteres, ~{len(texto_estudio)//4} tokens aprox)")
        except Exception as e:
            print(f"✗ ERROR al leer documento: {str(e)}")
            return JsonResponse({'error': f'Error al leer el documento: {str(e)}'}, status=400)
        
        # Extraer información usando Gemini
        try:
            print("✓ Llamando a Gemini API...")
            info_extraida = extraer_informacion_con_gemini(texto_estudio)
            print("\n" + "="*60)
            print("INFORMACIÓN EXTRAÍDA POR GEMINI:")
            print("="*60)
            print(f"Objeto: {info_extraida.get('objeto', 'NO EXTRAÍDO')}")
            print(f"Valor: {info_extraida.get('valor', 'NO EXTRAÍDO')}")
            print(f"Plazo: {info_extraida.get('plazo', 'NO EXTRAÍDO')}")
            print(f"Códigos UNSPSC: {info_extraida.get('codigos', 'NO EXTRAÍDO')}")
            print("="*60 + "\n")
        except Exception as e:
            print(f"✗ ERROR en Gemini: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'Error al procesar con Gemini: {str(e)}'}, status=500)
        
        # Obtener información del usuario
        user = request.user
        
        # Determinar género automáticamente (por defecto EL)
        # Si el usuario tiene first_name, intentar detectar género común
        genero = 'EL'  # Por defecto masculino
        if hasattr(user, 'first_name') and user.first_name:
            nombre = user.first_name.lower()
            # Nombres femeninos comunes terminan en 'a'
            if nombre.endswith('a') and not nombre.endswith('ia'):
                genero = 'LA'
        
        # Cargo por defecto
        cargo = 'PROFESIONAL UNIVERSITARIO'
        
        # Preparar datos para reemplazar (con límites de seguridad)
        MAX_OBJETO_CHARS = 2000  # Límite para el objeto del contrato
        MAX_CODIGOS_CHARS = 500  # Límite para códigos UNSPSC
        MAX_PLAZO_CHARS = 200    # Límite para plazo
        
        objeto_extraido = info_extraida.get('objeto', '')[:MAX_OBJETO_CHARS]
        codigos_extraidos = info_extraida.get('codigos', '')[:MAX_CODIGOS_CHARS]
        plazo_extraido = info_extraida.get('plazo', '')[:MAX_PLAZO_CHARS]
        
        datos = {
            'w_gen': genero,
            'w_cargo': cargo,
            'w_anno': str(datetime.now().year),
            'w_codigos': codigos_extraidos,
            'w_objeto': objeto_extraido,
            'w_valor': info_extraida.get('valor', ''),
            'w_plazo': plazo_extraido,
            'w_fecha': generar_fecha_en_letras()
        }
        
        print("\n" + "="*60)
        print("DATOS FINALES PARA EL CERTIFICADO:")
        print("="*60)
        print(f"Género: {datos['w_gen']}")
        print(f"Cargo: {datos['w_cargo']}")
        print(f"Año: {datos['w_anno']}")
        print(f"Códigos UNSPSC: {datos['w_codigos']}")
        print(f"Objeto: {datos['w_objeto'][:100]}..." if len(datos['w_objeto']) > 100 else f"Objeto: {datos['w_objeto']}")
        print(f"Valor: {datos['w_valor']}")
        print(f"Plazo: {datos['w_plazo']}")
        print(f"Fecha: {datos['w_fecha']}")
        print("="*60 + "\n")
        
        # Obtener formato solicitado
        formato = request.POST.get('formato', 'word')
        print(f"✓ Formato solicitado: {formato}")
        
        # Generar el certificado
        try:
            print("✓ Generando documento PAA...")
            certificado = generar_documento_paa(datos)
            print("✓ Documento generado exitosamente")
        except FileNotFoundError as e:
            print(f"✗ ERROR: Plantilla no encontrada - {str(e)}")
            return JsonResponse({'error': 'No se encontró la plantilla PAA. Por favor, coloque el archivo plantilla_paa.docx en la carpeta correcta.'}, status=500)
        except Exception as e:
            print(f"✗ ERROR al generar documento: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'Error al generar el certificado: {str(e)}'}, status=500)
        
        # Si se solicita PDF, convertir el documento
        if formato == 'pdf':
            if not DOCX2PDF_AVAILABLE:
                return JsonResponse({'error': 'La conversión a PDF no está disponible. Por favor, instale docx2pdf.'}, status=500)
            
            try:
                print("✓ Convirtiendo documento a PDF...")
                pdf_content = convertir_a_pdf(certificado)
                print("✓ Conversión a PDF exitosa")
                
                # Preparar respuesta PDF
                response = HttpResponse(
                    content=pdf_content,
                    content_type='application/pdf'
                )
                response['Content-Disposition'] = f'attachment; filename="Certificado_PAA_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
                
                print("✓ CERTIFICADO PDF GENERADO EXITOSAMENTE")
                print("="*60 + "\n")
                return response
            except Exception as e:
                print(f"✗ ERROR al convertir a PDF: {str(e)}")
                import traceback
                traceback.print_exc()
                return JsonResponse({'error': f'Error al convertir a PDF: {str(e)}'}, status=500)
        
        # Preparar respuesta Word
        print("✓ Preparando respuesta HTTP...")
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Certificado_PAA_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
        certificado.save(response)
        
        print("✓ CERTIFICADO GENERADO EXITOSAMENTE")
        print("="*60 + "\n")
        return response
        
    except Exception as e:
        # Log detallado del error
        print("\n" + "="*60)
        print("✗✗✗ ERROR GENERAL ✗✗✗")
        print("="*60)
        logger.error(f"Error en generar_certificado: {str(e)}", exc_info=True)
        import traceback
        error_detail = traceback.format_exc()
        print(f"{error_detail}")
        print("="*60 + "\n")
        return JsonResponse({'error': f'Error al procesar el documento: {str(e)}'}, status=500)


def extraer_informacion_con_gemini(texto_estudio):
    """Extrae información del estudio previo usando Gemini 2.5 Flash"""
    
    # API key hardcodeada temporalmente para pruebas
    api_key = "AIzaSyB-5c2Co7MuxWZnleGv_rjdSP7lrfOmRnM"
    
    print(f"✓ Configurando Gemini con API key: {api_key[:10]}...")
    genai.configure(api_key=api_key)
    
    # Usar el modelo Gemini 2.5 Pro (más potente y preciso)
    print("✓ Inicializando modelo gemini-2.5-pro...")
    model = genai.GenerativeModel('gemini-2.5-pro')
    
    # Prompt para extraer información
    prompt = f"""
Eres un asistente especializado en análisis de documentos de contratación pública.

Analiza el siguiente ESTUDIO PREVIO y extrae la información solicitada:

TEXTO DEL ESTUDIO PREVIO:
{texto_estudio}

INSTRUCCIONES:
Extrae la siguiente información del texto anterior:

1. **objeto**: Busca el "objeto del contrato", "descripción del bien o servicio", o similar. Extrae la descripción completa y conviértela a MAYÚSCULAS.

2. **valor**: Busca el "valor estimado", "presupuesto", "cuantía", o similar. Extrae SOLO el número sin símbolos de moneda ni puntos. Ejemplo: si dice "$1.234.567", extrae "1234567".

3. **plazo**: Busca el "plazo de ejecución", "duración del contrato", "tiempo de entrega", o similar. Extrae el texto tal como aparece (ejemplo: "12 MESES", "60 DÍAS").

4. **codigos**: CRÍTICO - Busca códigos de clasificación UNSPSC (Clasificador de Bienes y Servicios).

   INSTRUCCIONES DE BÚSQUEDA:
   - Busca números de 8 dígitos relacionados con UNSPSC
   - ACEPTA múltiples formatos:
     * Sin separadores: "72101500"
     * Con puntos: "72.10.15.00"
     * Con guiones: "72-10-15-00"
     * Con espacios: "72 10 15 00"
   
   ¿DÓNDE BUSCAR?
   - Busca en tablas con encabezados: "Código", "UNSPSC", "Clasificación", "Segmento", "Familia"
   - Busca cerca de descripciones de servicios o productos
   - Busca secciones con títulos como "Clasificación del bien o servicio según UNSPSC"
   
   LIMPIEZA DE DATOS:
   - Elimina puntos, espacios o guiones del resultado
   - Devuelve SOLO los 8 dígitos limpios (ejemplo: "72101500")
   - Si encuentras varios códigos, sepáralos por coma (ejemplo: "72101500, 43211500")
   
   IMPORTANTE:
   - IGNORA códigos CPC (suelen ser de 5 dígitos, ejemplo: "94900")
   - Si NO encuentras ningún código UNSPSC, deja este campo vacío ""

FORMATO DE RESPUESTA:
Responde ÚNICAMENTE con un objeto JSON válido, sin explicaciones adicionales:

{{
    "objeto": "DESCRIPCIÓN DEL OBJETO EN MAYÚSCULAS",
    "valor": "números con símbolos $ y separador de miles (punto)",
    "plazo": "PLAZO TAL COMO APARECE",
    "codigos": "código1, código2"
}}
"""
    
    # Generar respuesta (NO imprimir todo el documento en logs)
    print("\n" + "="*60)
    print("PREPARANDO ENVÍO A GEMINI")
    print("="*60)
    print(f"Total de caracteres a enviar: {len(texto_estudio)}")
    print(f"Tokens aproximados: ~{len(texto_estudio)//4}")
    print(f"Primeros 500 caracteres: {texto_estudio[:500]}...")
    print("="*60 + "\n")
    
    print("✓ Enviando prompt a Gemini...")
    try:
        response = model.generate_content(prompt)
    except Exception as e:
        error_msg = str(e)
        if "quota" in error_msg.lower() or "limit" in error_msg.lower():
            print(f"✗ Error de cuota/límite en Gemini: {error_msg}")
            raise Exception("Se excedió el límite de la API de Gemini. Por favor, intente con un documento más pequeño.")
        elif "token" in error_msg.lower():
            print(f"✗ Error de tokens en Gemini: {error_msg}")
            raise Exception("El documento es demasiado grande para procesar. Por favor, reduzca el tamaño del documento.")
        else:
            print(f"✗ Error al llamar a Gemini: {error_msg}")
            raise
    
    if not response or not response.text:
        print("✗ Gemini no devolvió respuesta válida")
        raise Exception("Gemini no devolvió respuesta válida")
    
    print(f"✓ Respuesta recibida de Gemini ({len(response.text)} caracteres)")
    texto_respuesta = response.text.strip()
    
    # Limpiar la respuesta para obtener solo el JSON
    # Remover markdown code blocks si existen
    texto_respuesta = re.sub(r'```json\s*', '', texto_respuesta)
    texto_respuesta = re.sub(r'```\s*', '', texto_respuesta)
    texto_respuesta = texto_respuesta.strip()
    
    # Buscar JSON entre llaves (igual que en el ejemplo)
    json_match = re.search(r'\{.*\}', texto_respuesta, re.DOTALL)
    
    # Parsear JSON
    import json
    try:
        if json_match:
            json_str = json_match.group()
            print(f"✓ JSON extraído: {json_str[:100]}...")
            info = json.loads(json_str)
        else:
            print("✗ No se encontró JSON válido en la respuesta")
            info = json.loads(texto_respuesta)
        
        print(f"✓ JSON parseado exitosamente: {info}")
        
    except json.JSONDecodeError as e:
        print(f"✗ Error parseando JSON: {str(e)}")
        print(f"Respuesta recibida: {texto_respuesta[:200]}...")
        # Si falla el parsing, usar valores por defecto
        info = {
            'objeto': 'NO IDENTIFICADO',
            'valor': '0',
            'plazo': 'NO ESPECIFICADO',
            'codigos': ''
        }
    
    return info


def generar_fecha_en_letras():
    """Genera la fecha actual en formato de letras en español"""
    meses = {
        1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
        5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
        9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
    }
    
    hoy = datetime.now()
    dia = hoy.day
    mes = meses[hoy.month]
    año = hoy.year
    
    # Convertir número a letras (simplificado para días 1-31)
    numeros = {
        1: 'uno', 2: 'dos', 3: 'tres', 4: 'cuatro', 5: 'cinco',
        6: 'seis', 7: 'siete', 8: 'ocho', 9: 'nueve', 10: 'diez',
        11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
        16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve',
        20: 'veinte', 21: 'veintiuno', 22: 'veintidós', 23: 'veintitrés',
        24: 'veinticuatro', 25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete',
        28: 'veintiocho', 29: 'veintinueve', 30: 'treinta', 31: 'treinta y uno'
    }
    
    dia_letras = numeros.get(dia, str(dia))
    
    return f"el {dia_letras} ({dia}) días del mes de {mes} de {año}"


def generar_documento_paa(datos):
    """Genera el documento PAA reemplazando los campos en la plantilla"""
    
    # Ruta a la plantilla
    base_dir = Path(__file__).resolve().parent
    plantilla_path = base_dir / 'templates' / 'paa' / 'plantilla_paa.docx'
    
    if not plantilla_path.exists():
        raise FileNotFoundError(f'No se encontró la plantilla en: {plantilla_path}')
    
    # Cargar la plantilla
    doc = Document(str(plantilla_path))
    
    # Reemplazar en párrafos
    for para in doc.paragraphs:
        for key, value in datos.items():
            placeholder = '{{' + key + '}}'
            if placeholder in para.text:
                # Reemplazar manteniendo el formato
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in datos.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in para.text:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
    
    # Reemplazar en encabezados y pies de página
    for section in doc.sections:
        # Encabezado
        header = section.header
        for para in header.paragraphs:
            for key, value in datos.items():
                placeholder = '{{' + key + '}}'
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
        
        # Pie de página
        footer = section.footer
        for para in footer.paragraphs:
            for key, value in datos.items():
                placeholder = '{{' + key + '}}'
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
    
    return doc


def convertir_a_pdf(documento_word):
    """
    Convierte un documento Word a PDF usando docx2pdf
    
    Args:
        documento_word: Objeto Document de python-docx
        
    Returns:
        bytes: Contenido del PDF generado
    """
    # Crear archivos temporales
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
        temp_docx_path = temp_docx.name
        documento_word.save(temp_docx_path)
    
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
        temp_pdf_path = temp_pdf.name
    
    try:
        # Convertir DOCX a PDF
        print(f"✓ Convirtiendo {temp_docx_path} a {temp_pdf_path}")
        convert(temp_docx_path, temp_pdf_path)
        
        # Leer el contenido del PDF
        with open(temp_pdf_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()
        
        return pdf_content
        
    finally:
        # Limpiar archivos temporales
        try:
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
        except Exception as e:
            print(f"⚠️ Error al limpiar archivos temporales: {str(e)}")
