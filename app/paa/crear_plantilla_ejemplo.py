"""
Script para crear una plantilla PAA de ejemplo.
Ejecutar desde la carpeta app/paa:
    python crear_plantilla_ejemplo.py
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def crear_plantilla_paa():
    """Crea una plantilla PAA de ejemplo con formato básico"""
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Encabezado
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "EMPRESA DE DISTRIBUCIÓN ELÉCTRICA DEL NORTE S.A. - EDENORTE"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.runs[0].font.size = Pt(10)
    header_para.runs[0].font.bold = True
    
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("CERTIFICADO DEL PLAN ANUAL DE ADQUISICIONES")
    run_titulo.font.size = Pt(14)
    run_titulo.font.bold = True
    
    doc.add_paragraph()  # Espacio
    
    # Contenido principal
    contenido = doc.add_paragraph()
    contenido.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    texto = (
        "{{w_gen}} SUSCRITO {{w_cargo}} DE LA EMPRESA DE DISTRIBUCIÓN ELÉCTRICA DEL NORTE S.A. "
        "(EDENORTE), CERTIFICA QUE EN EL PLAN ANUAL DE ADQUISICIONES DEL AÑO {{w_anno}}, "
        "SE ENCUENTRA INCLUIDO EL SIGUIENTE PROCESO CONTRACTUAL:"
    )
    
    run_contenido = contenido.add_run(texto)
    run_contenido.font.size = Pt(12)
    
    doc.add_paragraph()  # Espacio
    
    # Códigos UNSPSC
    p_codigos = doc.add_paragraph()
    p_codigos.add_run("CÓDIGOS UNSPSC: ").font.bold = True
    p_codigos.add_run("{{w_codigos}}")
    p_codigos.runs[0].font.size = Pt(12)
    p_codigos.runs[1].font.size = Pt(12)
    
    doc.add_paragraph()  # Espacio
    
    # Objeto
    p_objeto = doc.add_paragraph()
    p_objeto.add_run("OBJETO: ").font.bold = True
    p_objeto.add_run("{{w_objeto}}")
    p_objeto.runs[0].font.size = Pt(12)
    p_objeto.runs[1].font.size = Pt(12)
    p_objeto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()  # Espacio
    
    # Valor
    p_valor = doc.add_paragraph()
    p_valor.add_run("VALOR ESTIMADO: ").font.bold = True
    p_valor.add_run("${{w_valor}} PESOS")
    p_valor.runs[0].font.size = Pt(12)
    p_valor.runs[1].font.size = Pt(12)
    
    doc.add_paragraph()  # Espacio
    
    # Plazo
    p_plazo = doc.add_paragraph()
    p_plazo.add_run("PLAZO: ").font.bold = True
    p_plazo.add_run("{{w_plazo}}")
    p_plazo.runs[0].font.size = Pt(12)
    p_plazo.runs[1].font.size = Pt(12)
    
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    
    # Fecha
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    texto_fecha = "SE EXPIDE LA PRESENTE CERTIFICACIÓN A LOS {{w_fecha}}."
    run_fecha = p_fecha.add_run(texto_fecha)
    run_fecha.font.size = Pt(12)
    
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    
    # Firma
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.add_run("_" * 50)
    
    p_nombre = doc.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nombre = p_nombre.add_run("{{w_cargo}}")
    run_nombre.font.size = Pt(11)
    run_nombre.font.bold = True
    
    p_empresa = doc.add_paragraph()
    p_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_empresa = p_empresa.add_run("EDENORTE")
    run_empresa.font.size = Pt(11)
    
    # Pie de página
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Página 1 de 1"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    
    # Guardar documento
    output_path = Path(__file__).parent / 'templates' / 'paa' / 'plantilla_paa.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    print(f"✅ Plantilla PAA creada exitosamente en: {output_path}")
    print("\n📝 Placeholders incluidos:")
    print("   - {{w_gen}}")
    print("   - {{w_cargo}}")
    print("   - {{w_anno}}")
    print("   - {{w_codigos}}")
    print("   - {{w_objeto}}")
    print("   - {{w_valor}}")
    print("   - {{w_plazo}}")
    print("   - {{w_fecha}}")
    print("\n⚠️  Recuerde personalizar esta plantilla según el formato oficial de EDENORTE")

if __name__ == '__main__':
    crear_plantilla_paa()
